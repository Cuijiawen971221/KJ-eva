import docx
import csv
import argparse
from typing import List, Dict

def extract_paragraphs_by_style(doc_path: str, target_style: str) -> List[str]:
    """按样式提取Word文档中的段落"""
    doc = docx.Document(doc_path)
    return [para.text for para in doc.paragraphs if para.style.name == target_style]




def extract_structured_data(doc_path: str) -> List[Dict[str, str]]:
    """从Word文档中提取结构化数据（修复空行判断和记录分割逻辑）"""
    doc = docx.Document(doc_path)
    data = []
    current_item = {}
    # 核心字段：出现时代表新记录开始（如station_code是每条气象记录的唯一标识）
    core_fields = {"station_code"}  # 针对气象数据的核心字段，确保新记录触发
    
    for para in doc.paragraphs:
        raw_text = para.text  # 保留原始文本（用于判断空行）
        text = raw_text.strip()  # 去除前后空白（用于提取键值对）
        
        # 1. 处理空行/伪空行：触发当前记录保存
        if not text:  # 不管原始文本是否含空白字符，只要strip后为空就视为空行
            if current_item:  # 若当前有未保存的记录
                data.append(current_item)
                current_item = {}  # 重置，准备接收新记录
            continue  # 跳过空行，处理下一段
        
        # 2. 处理键值对（支持英文冒号:）
        if ":" in text:
            separator = ":"
            key_value = text.split(separator, 1)  # 只分割第一个冒号（避免值中含冒号）
            if len(key_value) == 2:
                key = key_value[0].strip()
                value = key_value[1].strip()
                
                # 3. 兜底逻辑：若遇到核心字段（如station_code），说明是新记录
                if key in core_fields and current_item:
                    data.append(current_item)  # 先保存上一条记录
                    current_item = {}  # 重置为新记录
                    
                current_item[key] = value
    
    # 4. 保存最后一条未被空行触发的记录
    if current_item:
        data.append(current_item)
        
    return data


def extract_table_data(doc_path: str) -> List[Dict[str, str]]:
    """从Word文档的表格中提取数据"""
    doc = docx.Document(doc_path)
    all_data = []
    
    for table in doc.tables:
        # 假设第一行为表头
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # 提取表格内容
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            # 确保表头和内容长度一致
            if len(row_data) == len(headers):
                all_data.append(dict(zip(headers, row_data)))
    
    return all_data

def save_to_csv(data: List[Dict[str, str]], csv_path: str = "output.csv") -> None:
    """将数据保存为CSV文件，默认输出为output.csv"""
    if not data:
        print("没有提取到数据，不生成CSV文件")
        return
        
    # 获取所有字段名（取并集）
    fieldnames = set()
    for item in data:
        fieldnames.update(item.keys())
    fieldnames = sorted(fieldnames)
    
    with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(data)
    
    print(f"数据已成功保存到 {csv_path}，共 {len(data)} 条记录")

def decode(input_path, mode):
    """
    解码Word文档并转换为CSV
    
    Args:
        input_path: Word文档路径（.docx格式）
        mode: 提取模式，'structure'（结构化文本）或 'table'（表格）
    """
    try:
        if mode == 'structure':
            data = extract_structured_data(input_path)
        elif mode == 'table':
            data = extract_table_data(input_path)
        else:
            raise ValueError(f"不支持的模式: {mode}，请使用 'structure' 或 'table'")
            
        # 输出固定为output.csv
        save_to_csv(data)
        
    except Exception as e:
        print(f"转换过程中发生错误：{str(e)}")

if __name__ == "__main__":
    
    decode("demo.docx", "structure")
    
